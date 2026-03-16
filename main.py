import os
import re
import datetime
import logging
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.scrollview import ScrollView
from kivy.uix.progressbar import ProgressBar
from kivy.uix.popup import Popup
from kivy.core.window import Window
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor

# 设置 Android 字体支持（解决中文乱码）
from kivy.core.text import LabelBase
# 如果你的安卓系统有默认中文字体路径，可以在此注册，否则 Kivy 会尝试使用系统默认

class QuestionExtractorApp(App):
    def build(self):
        self.title = "试题提取工具"
        self.questions_data = []
        self.stats = {"单选题": 0, "多选题": 0, "判断题": 0, "总数": 0}
        self.current_filename = "未命名"
        
        # 安卓存储路径
        self.base_path = "/storage/emulated/0/试题提取工具"
        self.export_dir = os.path.join(self.base_path, "导出文件")
        
        # 主布局
        root = BoxLayout(orientation='vertical', spacing=10, padding=10)
        
        # 头部
        header = Label(text="试题提取工具", size_hint_y=None, height=50, 
                      color=(1, 1, 1, 1), font_size='20sp', bold=True)
        root.add_widget(header)

        # 按钮区
        btn_layout = BoxLayout(orientation='vertical', spacing=5, size_hint_y=None, height=350)
        
        btns = [
            ("选择 Excel 文件 (需放在根目录)", self.open_file_mock), # 安卓下建议直接读取特定路径
            ("导出标准 DOCX", lambda x: self.export_data("standard")),
            ("导出手机阅读版", lambda x: self.export_data("phone")),
            ("导出纯文本 TXT", lambda x: self.export_data("txt")),
            ("统计结果", self.show_stats),
            ("关于", self.show_about)
        ]

        for text, cmd in btns:
            btn = Button(text=text, size_hint_y=None, height=50, 
                         background_color=(0.9, 0, 0.07, 1)) # 吉林银行红
            btn.bind(on_release=cmd)
            btn_layout.add_widget(btn)
        
        root.add_widget(btn_layout)

        # 进度条
        self.pb = ProgressBar(max=100, size_hint_y=None, height=20)
        root.add_widget(self.pb)

        # 日志显示区
        self.scroll = ScrollView()
        self.log_label = Label(text="等待导入文件...\n注意：请将 Excel 放入手机存储的“试题提取工具”文件夹中", 
                               size_hint_y=None, halign='left', valign='top',
                               text_size=(Window.width - 20, None))
        self.log_label.bind(texture_size=self.log_label.setter('text_size'))
        self.scroll.add_widget(self.log_label)
        root.add_widget(self.scroll)

        return root

    def clean_text(self, text):
        if text is None: return ""
        text = str(text).replace('\n', '').replace('\r', '').strip()
        text = re.sub(r'\(+\s*\)+', '(　)', text)
        text = re.sub(r'（+\s*）+', '（　）', text)
        return text

    def open_file_mock(self, instance):
        # 由于安卓文件选择器调用复杂，这里演示读取 /storage/emulated/0/试题提取工具/input.xlsx
        target_file = os.path.join(self.base_path, "input.xlsx")
        if os.path.exists(target_file):
            self.process_excel(target_file)
        else:
            self.show_popup("找不到文件", f"请将文件重命名为 input.xlsx 并放入 {self.base_path}")

    def process_excel(self, file_path):
        try:
            wb = load_workbook(file_path, data_only=True)
            self.questions_data = []
            self.stats = {"单选题": 0, "多选题": 0, "判断题": 0, "总数": 0}
            
            for sheet in wb.worksheets:
                rows = list(sheet.rows)
                if not rows: continue
                title_idx, ans_idx = -1, -1
                header_row_idx = -1
                for i, row in enumerate(rows):
                    vals = [str(c.value) if c.value else "" for c in row]
                    if any("题目" in v for v in vals):
                        header_row_idx = i
                        for j, v in enumerate(vals):
                            if "题目" in v: title_idx = j
                            if "正确答案" in v: ans_idx = j
                        break
                
                if title_idx == -1: continue
                
                for r in range(header_row_idx + 1, len(rows)):
                    row = rows[r]
                    if not row[title_idx].value: continue
                    title = self.clean_text(row[title_idx].value)
                    ans_val = self.clean_text(row[ans_idx].value).upper()
                    
                    options = []
                    for c in range(ans_idx + 1, len(row)):
                        opt_val = self.clean_text(row[c].value)
                        if opt_val:
                            options.append(opt_val)
                    
                    q_type = "单选题"
                    if ans_val in ["对", "错", "正确", "错误"]: q_type = "判断题"
                    elif "," in ans_val or len(ans_val) > 1: q_type = "多选题"
                    
                    self.questions_data.append({"title": title, "ans": ans_val, "opts": options, "type": q_type})
                    self.stats[q_type] += 1
                    self.stats["总数"] += 1

            self.log_label.text = f"成功解析 {self.stats['总数']} 道题！可以开始导出了。"
        except Exception as e:
            self.show_popup("错误", str(e))

    def export_data(self, mode):
        if not self.questions_data:
            return self.show_popup("提示", "请先打开并解析 Excel 文件")
        
        if not os.path.exists(self.export_dir): os.makedirs(self.export_dir)
        path = os.path.join(self.export_dir, f"export_{mode}.docx")
        
        try:
            doc = Document()
            for i, q in enumerate(self.questions_data):
                p = doc.add_paragraph(f"{i+1}. {q['title']}")
                for o in q['opts']:
                    doc.add_paragraph(f"   {o}")
                doc.add_paragraph(f"答案：{q['ans']}")
            doc.save(path)
            self.show_popup("成功", f"文件已保存至：\n{path}")
        except Exception as e:
            self.show_popup("导出失败", str(e))

    def show_stats(self, instance):
        msg = f"总数：{self.stats['总数']}\n单选：{self.stats['单选题']}\n多选：{self.stats['多选题']}\n判断：{self.stats['判断题']}"
        self.show_popup("统计结果", msg)

    def show_about(self, instance):
        self.show_popup("关于", "试题提取工具安卓版\n作者：🍍")

    def show_popup(self, title, content):
        p = Popup(title=title, content=Label(text=content), size_hint=(0.8, 0.5))
        p.open()

if __name__ == "__main__":
    QuestionExtractorApp().run()